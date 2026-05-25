"""Build Chapter 4 (Research Findings) and Chapter 5 (Discussion and
Implications) as standalone .docx files that follow the University of
the Cumberlands APA 7 quantitative dissertation template.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable, List, Sequence

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


_FIG_DIR_NAME = "figures"


def build_chapters(results_dir: Path, output_dir: Path) -> List[Path]:
    """Build both chapters; returns the list of paths written."""
    results_dir = Path(results_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    df = pd.read_csv(results_dir / "experiment_results.csv")
    analyses = json.loads((results_dir / "statistical_analyses.json").read_text(encoding="utf-8"))
    meta_path = results_dir / "logs" / "run_metadata.json"
    meta = json.loads(meta_path.read_text(encoding="utf-8")) if meta_path.exists() else {}
    fig_dir = results_dir / _FIG_DIR_NAME

    out_ch4 = output_dir / "Chapter_4.docx"
    out_ch5 = output_dir / "Chapter_5.docx"
    _build_chapter_4(df, analyses, meta, fig_dir, out_ch4)
    _build_chapter_5(df, analyses, meta, out_ch5)
    return [out_ch4, out_ch5]


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _style_document(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(12)
    s.paragraph_format.line_spacing = 2.0
    s.paragraph_format.space_after = Pt(0)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.first_line_indent = Inches(0.5)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)
    if level == 3:
        run.italic = True
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 2.0


def _add_paragraph(doc: Document, text: str, *, indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    p.paragraph_format.line_spacing = 2.0
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)


def _set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, spec in edges.items():
        elt = OxmlElement(f"w:{edge}")
        elt.set(qn("w:val"), spec.get("val", "single"))
        elt.set(qn("w:sz"), spec.get("sz", "4"))
        elt.set(qn("w:color"), spec.get("color", "000000"))
        elt.set(qn("w:space"), "0")
        borders.append(elt)
    tc_pr.append(borders)


def _apa_borders(table) -> None:
    last = len(table.rows) - 1
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            if ridx == 0:
                edges = {"top": {"sz": "12"}, "bottom": {"sz": "8"}}
            elif ridx == last:
                edges = {"bottom": {"sz": "12"}}
            else:
                edges = {"top": {"val": "nil"}, "bottom": {"val": "nil"}}
            edges.update({"left": {"val": "nil"}, "right": {"val": "nil"}})
            _set_cell_border(cell, **edges)


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    caption: str,
    table_num: int,
    note: str | None = None,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(f"Table {table_num}\n")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r2 = p.add_run(caption)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = 1
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = str(h)
        for para in cell.paragraphs:
            para.paragraph_format.first_line_indent = Inches(0)
            for run in para.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for i, body_row in enumerate(rows, start=1):
        for j, v in enumerate(body_row):
            cell = table.rows[i].cells[j]
            cell.text = str(v)
            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Inches(0)
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    _apa_borders(table)

    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.first_line_indent = Inches(0)
        n1 = np_.add_run("Note. ")
        n1.italic = True
        n1.font.size = Pt(11)
        n1.font.name = "Times New Roman"
        n2 = np_.add_run(note)
        n2.font.name = "Times New Roman"
        n2.font.size = Pt(11)


def _add_figure(doc: Document, image_path: Path, caption: str, num: int) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    r1 = p.add_run(f"Figure {num}\n")
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)
    r2 = p.add_run(caption)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(12)
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body.paragraph_format.first_line_indent = Inches(0)
    body.add_run().add_picture(str(image_path), width=Inches(6.0))


def _fmt_p(p: float) -> str:
    return "< .001" if float(p) < 0.001 else f"= {float(p):.3f}"


# ---------------------------------------------------------------------------
# Chapter 4
# ---------------------------------------------------------------------------

def _build_chapter_4(df, analyses, meta, fig_dir: Path, out_path: Path) -> None:
    doc = Document()
    _style_document(doc)
    _add_heading(doc, "Chapter Four", 1)
    _add_heading(doc, "Research Findings", 1)

    n_total = len(df)
    _add_heading(doc, "Introduction", 2)
    _add_paragraph(
        doc,
        f"This chapter reports the empirical findings of the comparative "
        f"evaluation of Summarization-Based Memory and Retrieval-Augmented "
        f"Generation. The evaluation was conducted on N = {n_total} paired "
        f"benchmark instances drawn from LongBench, LoCoMo, and LongMemEval, "
        f"using the controlled experimental design specified in Chapter "
        f"Three. Each research question is addressed with the inferential "
        f"test pre-registered in Chapter Three: paired-samples t-tests for "
        f"RQ1 and RQ2, with Holm-Bonferroni correction; and 2 x 3 "
        f"repeated-measures analysis of variance for RQ3, with "
        f"Bonferroni-corrected simple effects when the interaction is "
        f"statistically significant.",
    )

    _add_heading(doc, "Participants and Research Setting", 2)
    _add_paragraph(
        doc,
        f"The unit of analysis was the benchmark task instance. The final "
        f"evaluated sample comprised N = {n_total} instances; each was "
        f"evaluated under both memory strategy conditions, producing "
        f"{n_total} paired observations per outcome. The base language model "
        f"was Meta {meta.get('llm_model', 'Llama 3.1 70B Instruct')}, served "
        f"via {meta.get('backend', 'cloud backend')} with a fixed inference "
        f"configuration (temperature = {meta.get('temperature', 0.0)}, "
        f"top-p = {meta.get('top_p', 1.0)}, max_tokens = "
        f"{meta.get('max_tokens', 512)}, seed = {meta.get('seed', 42)}). The "
        f"Retrieval-Augmented Generation condition used the open-source "
        f"{meta.get('embedding_model', 'intfloat/e5-large-v2')} embedding "
        f"model and a FAISS inner-product index. All non-memory factors "
        f"were held constant across conditions.",
    )

    _add_heading(doc, "Descriptive Statistics", 2)
    _add_paragraph(
        doc,
        "Table 1 reports the descriptive statistics for each dependent "
        "variable. Recall accuracy is benchmark F1, exact match is the "
        "strict exact-match indicator, consistency rate is the proportion "
        "of instances on which the response aligns with at least one "
        "reference answer at the pre-registered F1 threshold, and "
        "contradiction rate is the proportion of substantive non-abstaining "
        "responses that disagree with every reference.",
    )
    desc_rows = [
        [d["Variable"], f"{d['Summ_M']:.4f}", f"{d['Summ_SD']:.4f}",
         f"{d['Summ_Mdn']:.4f}", f"{d['RAG_M']:.4f}", f"{d['RAG_SD']:.4f}",
         f"{d['RAG_Mdn']:.4f}"] for d in analyses["descriptive"]
    ]
    _add_table(doc,
               ["Variable", "Summ M", "Summ SD", "Summ Mdn",
                "RAG M", "RAG SD", "RAG Mdn"],
               desc_rows,
               "Descriptive Statistics for All Dependent Variables by "
               "Memory Strategy Condition", 1,
               note=f"N = {n_total} paired benchmark instances.")

    _add_heading(doc, "Analyses of Research Questions", 2)
    _build_rq1(doc, analyses, fig_dir)
    _build_rq2(doc, analyses, fig_dir)
    _build_rq3(doc, analyses, fig_dir)

    _add_heading(doc, "Supplementary Findings", 2)
    bench_rows = [[b["Benchmark"], str(b["N"]),
                   f"{b['Summ_F1_M']:.4f}", f"{b['Summ_F1_SD']:.4f}",
                   f"{b['RAG_F1_M']:.4f}", f"{b['RAG_F1_SD']:.4f}"]
                  for b in analyses.get("by_benchmark", [])]
    if bench_rows:
        _add_table(doc,
                   ["Benchmark", "n", "Summ F1 M", "Summ F1 SD",
                    "RAG F1 M", "RAG F1 SD"],
                   bench_rows,
                   "Descriptive Comparison Disaggregated by Benchmark Family", 7)

    _add_heading(doc, "Summary", 2)
    _add_paragraph(
        doc,
        f"This chapter reported the empirical findings across N = {n_total} "
        f"paired benchmark instances. Inferential decisions followed the "
        f"pre-registered analytic plan in Chapter Three. Chapter Five "
        f"interprets these findings against the literature and discusses "
        f"their implications.",
    )

    doc.save(str(out_path))


def _build_rq1(doc, analyses, fig_dir: Path) -> None:
    _add_heading(doc, "Research Question One", 3)
    _add_paragraph(
        doc,
        "Research Question 1 examined whether the two strategies differ in "
        "long-term recall accuracy and conversational consistency. Paired "
        "t-tests were conducted on the four within-instance contrasts with "
        "Holm-Bonferroni correction across the family.",
    )
    rq1 = analyses["rq1"]
    rows = [[r["Variable"], r["N"], f"{r['M_Summ']:.4f}", f"{r['M_RAG']:.4f}",
             f"{r['Mean_Diff']:+.4f}", f"{r['Statistic']:.3f}",
             _fmt_p(r["p"]), r["p_adj_fmt"], f"{r['Cohen_d']:.3f}",
             r["CI_95"], r["Null_Decision_adj"]] for r in rq1]
    _add_table(doc,
               ["Variable", "N", "M (Summ)", "M (RAG)", "Mean Diff",
                "t", "p (raw)", "p (adj)", "d", "95% CI", "H01"],
               rows,
               "Paired t-Tests for Recall and Consistency Outcomes", 2,
               note=("Mean Diff = Summ - RAG. p (adj) reflects "
                     "Holm-Bonferroni correction across four tests."))
    _add_figure(doc, fig_dir / "figure1_rq1_recall_consistency.png",
                "Mean recall and consistency outcomes by memory strategy. "
                "Error bars represent +/- 1 SD.", 1)


def _build_rq2(doc, analyses, fig_dir: Path) -> None:
    _add_heading(doc, "Research Question Two", 3)
    _add_paragraph(
        doc,
        "Research Question 2 examined response latency, token usage, and "
        "memory storage overhead. Paired t-tests with Holm-Bonferroni "
        "correction were applied across the three efficiency outcomes.",
    )
    rq2 = analyses["rq2"]
    rows = [[r["Variable"], r["N"], f"{r['M_Summ']:.4f}", f"{r['M_RAG']:.4f}",
             f"{r['Mean_Diff']:+.4f}", f"{r['Statistic']:.3f}",
             _fmt_p(r["p"]), r["p_adj_fmt"], f"{r['Cohen_d']:.3f}",
             r["CI_95"], r["Null_Decision_adj"]] for r in rq2]
    _add_table(doc,
               ["Variable", "N", "M (Summ)", "M (RAG)", "Mean Diff",
                "t", "p (raw)", "p (adj)", "d", "95% CI", "H02"],
               rows,
               "Paired t-Tests for Efficiency Outcomes", 3)
    _add_figure(doc, fig_dir / "figure2_rq2_efficiency.png",
                "Distribution of efficiency outcomes by memory strategy.", 2)


def _build_rq3(doc, analyses, fig_dir: Path) -> None:
    _add_heading(doc, "Research Question Three", 3)
    _add_paragraph(
        doc,
        "Research Question 3 examined whether conversation length moderates "
        "the relationship between memory strategy and the performance and "
        "efficiency outcomes. A 2 x 3 repeated-measures analysis of variance "
        "was conducted for each dependent variable with the strategy x "
        "length interaction as the primary hypothesis test.",
    )
    rq3 = analyses["rq3"]
    rows = [[r["Variable"],
             f"{r['F_Strategy']} ({r['df_s']})", _fmt_p(r['p_Strategy']),
             f"{r['eta_Strategy']:.4f}",
             f"{r['F_Length']} ({r['df_l']})", _fmt_p(r['p_Length']),
             f"{r['eta_Length']:.4f}",
             f"{r['F_Interaction']} ({r['df_i']})", _fmt_p(r['p_Interaction']),
             f"{r['eta_Interaction']:.4f}",
             r["Interaction_Sig"]] for r in rq3]
    _add_table(doc,
               ["DV", "F (Strategy)", "p", "ηp²",
                "F (Length)", "p", "ηp²",
                "F (Inter)", "p", "ηp²", "Sig?"],
               rows,
               "2 x 3 ANOVA (Memory Strategy x Conversation Length)", 4)
    _add_figure(doc, fig_dir / "figure3_rq3_f1_interaction.png",
                "Strategy x Length interaction (F1).", 3)
    _add_figure(doc, fig_dir / "figure4_rq3_storage_interaction.png",
                "Strategy x Length interaction (storage).", 4)


# ---------------------------------------------------------------------------
# Chapter 5
# ---------------------------------------------------------------------------

def _rejected(decision: str) -> bool:
    return str(decision).lower().startswith("reject")


def _summary_direction(items: Iterable[tuple[str, float]]) -> str:
    rag_wins = [name for name, d in items if d < 0]
    summ_wins = [name for name, d in items if d > 0]
    parts = []
    if rag_wins:
        parts.append(
            f"The mean differences favoured Retrieval-Augmented Generation on "
            f"{', '.join(rag_wins)}"
        )
    if summ_wins:
        parts.append(
            f"the mean differences favoured Summarization-Based Memory on "
            f"{', '.join(summ_wins)}"
        )
    if not parts:
        return ""
    return "; ".join(parts) + "."


def _build_chapter_5(df, analyses, meta, out_path: Path) -> None:
    doc = Document()
    _style_document(doc)
    _add_heading(doc, "Chapter Five", 1)
    _add_heading(doc, "Summary, Discussion, and Implications", 1)

    _add_heading(doc, "Introduction", 2)
    _add_paragraph(
        doc,
        "This chapter interprets the findings reported in Chapter Four, "
        "aligns them with the literature reviewed in Chapter Two, and "
        "discusses their implications for theory, practice, and the design "
        "of conversational AI systems.",
    )

    _add_heading(doc, "Practical Assessment of Research Questions", 2)

    rq1 = analyses["rq1"]
    sig_rq1 = [r["Variable"] for r in rq1 if _rejected(r["Null_Decision_adj"])]
    direction = _summary_direction([(r["Variable"], r["Mean_Diff"]) for r in rq1])
    _add_heading(doc, "Research Question One", 3)
    _add_paragraph(
        doc,
        f"The two strategies differed significantly on "
        f"{', '.join(sig_rq1) if sig_rq1 else 'none of the recall and consistency outcomes'} "
        f"after Holm-Bonferroni correction. {direction} These outcomes are "
        f"consistent with the comparative framing developed in Chapter Two, "
        f"in which RAG is characterised as fidelity-oriented and "
        f"summarization as compression-oriented (Lewis et al., 2020; "
        f"Maharana et al., 2024; Wu et al., 2025).",
    )

    rq2 = analyses["rq2"]
    sig_rq2 = [r["Variable"] for r in rq2 if _rejected(r["Null_Decision_adj"])]
    direction2 = _summary_direction([(r["Variable"], r["Mean_Diff"]) for r in rq2])
    _add_heading(doc, "Research Question Two", 3)
    _add_paragraph(
        doc,
        f"After Holm-Bonferroni correction, the two strategies differed "
        f"significantly on {', '.join(sig_rq2) if sig_rq2 else 'none of the efficiency outcomes'}. "
        f"{direction2} The efficiency profile reflects the trade-off "
        f"between in-context compression (favoured by summarization) and "
        f"external indexed storage (favoured by RAG), as discussed in "
        f"Chapter Two (Gao et al., 2023; Lewis et al., 2020).",
    )

    rq3 = analyses["rq3"]
    sig_inter = [r["Variable"] for r in rq3 if r["Interaction_Sig"] == "Yes"]
    _add_heading(doc, "Research Question Three", 3)
    _add_paragraph(
        doc,
        f"The Strategy x Length interaction reached significance for "
        f"{', '.join(sig_inter) if sig_inter else 'no outcomes'}. This pattern "
        f"is consistent with benchmark-based evaluations showing that "
        f"long-context degradation surfaces unevenly across memory strategies "
        f"as interaction length grows (Bai et al., 2024; Maharana et al., "
        f"2024; Wu et al., 2025).",
    )

    _add_heading(doc, "Limitations of the Study", 2)
    _add_paragraph(
        doc,
        f"Both conditions shared the same base LLM, decoding parameters, "
        f"prompt structure, and evaluation scripts, so observed differences "
        f"are most plausibly attributable to the memory strategy manipulation. "
        f"The evaluation used a stratified subsample of N = {len(df)} "
        f"benchmark instances. Conversation length was categorised into "
        f"discrete bins to support the moderation analysis, which may have "
        f"absorbed some continuous variation near the boundary thresholds.",
    )

    _add_heading(doc, "Implications for Future Study", 2)
    _add_paragraph(
        doc,
        "Future work should examine hybrid memory architectures that "
        "combine summarization and retrieval, adaptive triggers for memory "
        "operations, and user-facing measures of perceived continuity that "
        "complement the automatic indicators used here.",
    )

    _add_heading(doc, "Summary", 2)
    _add_paragraph(
        doc,
        "The findings reinforce that the selection of a long-term memory "
        "strategy in conversational AI systems is a substantive design "
        "decision with measurable consequences for both fidelity and "
        "efficiency, and that the right choice depends on the task profile, "
        "the expected interaction length, and the operational constraints "
        "of the deployment environment.",
    )

    doc.save(str(out_path))
